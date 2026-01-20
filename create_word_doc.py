"""
Convert Technical PRD Markdown to Word Document
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import re
from datetime import datetime

# Create document
doc = Document()

# Set document properties
doc.core_properties.title = "Technical PRD: Support-Led Ordering System"
doc.core_properties.author = "Engineering Team"

# Styles
title_style = doc.styles['Title']
title_style.font.size = Pt(28)
title_style.font.bold = True

heading1_style = doc.styles['Heading 1']
heading1_style.font.size = Pt(18)
heading1_style.font.bold = True
heading1_style.font.color.rgb = RGBColor(0, 51, 102)

heading2_style = doc.styles['Heading 2']
heading2_style.font.size = Pt(14)
heading2_style.font.bold = True

# Title Page
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph("Technical PRD")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(36)
title.runs[0].font.bold = True

subtitle = doc.add_paragraph("Support-Led Ordering System")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(24)

doc.add_paragraph()
subtitle2 = doc.add_paragraph("Customer Support Backend with Freshchat Integration")
subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle2.runs[0].font.size = Pt(16)
subtitle2.runs[0].font.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Metadata table
meta_table = doc.add_table(rows=4, cols=2)
meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
meta_data = [
    ("Version", "1.0"),
    ("Date", datetime.now().strftime("%B %d, %Y")),
    ("Status", "Draft"),
    ("Author", "Engineering Team")
]
for i, (key, value) in enumerate(meta_data):
    meta_table.rows[i].cells[0].text = key
    meta_table.rows[i].cells[1].text = value

doc.add_page_break()

# Table of Contents placeholder
doc.add_heading("Table of Contents", level=1)
toc_items = [
    "1. Overview",
    "2. User Flow",
    "3. System Architecture",
    "4. Data Models",
    "5. API Specifications",
    "6. Service Layer",
    "7. Mobile Integration",
    "8. Analytics & Reporting",
    "9. Deployment",
    "10. Implementation Timeline",
    "11. Success Metrics"
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')

doc.add_page_break()

# 1. Overview
doc.add_heading("1. Overview", level=1)

doc.add_heading("1.1 Problem Statement", level=2)
doc.add_paragraph(
    "Current catering order flow relies heavily on human sales intervention. "
    "Users abandon orders due to trust issues, menu/quantity confusion, and high cognitive load with choices."
)

doc.add_heading("1.2 Solution", level=2)
doc.add_paragraph("Build a Support-Led Ordering System where:")
solution_items = [
    "App is the primary interface",
    "Human support is contextual (triggered only when needed)",
    "Every interaction is logged as an \"Incidence\" for product improvement"
]
for item in solution_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("1.3 Selected Platform", level=2)
doc.add_paragraph("Freshchat Growth — chosen based on:")
platform_items = [
    "Native React Native SDK support",
    "Budget fit (₹9,600/mo for 6 agents)",
    "Robust webhooks and REST API",
    "2-week implementation timeline"
]
for item in platform_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 2. User Flow (NEW SECTION)
doc.add_heading("2. User Flow", level=1)

# Add the user flow diagram image
doc.add_paragraph("Visual Overview:", style='Intense Quote')
try:
    doc.add_picture(r"C:\Users\Syed Ashfaque Hussai\OneDrive\Desktop\Customer Support\User_Flow_Diagram.png", width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
except Exception as e:
    doc.add_paragraph(f"[User Flow Diagram - Image not found: {e}]")

doc.add_paragraph()

doc.add_heading("2.1 End-to-End Support Flow", level=2)
doc.add_paragraph(
    "This section describes the complete user journey from browsing the app to getting support assistance."
)


# Step 1
doc.add_heading("Step 1: User Browses the App", level=3)
doc.add_paragraph(
    "The user opens the mobile app (React Native) and browses menus, customizes platters, adds items to cart."
)
doc.add_paragraph("Behind the scenes:", style='Intense Quote')
step1_items = [
    "Context Tracker monitors: current screen, cart updates, navigation patterns",
    "This data is continuously synced to Freshchat via setUserProperties()"
]
for item in step1_items:
    doc.add_paragraph(item, style='List Bullet')

# Step 2
doc.add_heading("Step 2: User Requests Support", level=3)
doc.add_paragraph("User clicks \"Chat with Support\" or a help button.")
doc.add_paragraph("What happens:", style='Intense Quote')
step2_items = [
    "Frontend sends current context to Backend (POST /context/update)",
    "Context includes: current_screen, cart_value, guest_count, behavior signals",
    "Example: cart_value=₹25,000, guest_count=100, inactivity_seconds=75"
]
for item in step2_items:
    doc.add_paragraph(item, style='List Bullet')

# Step 3
doc.add_heading("Step 3: Backend Processes the Request", level=3)
doc.add_paragraph("The backend performs three key operations:")

doc.add_paragraph("A. Friction Score Calculation:", style='Intense Quote')
friction_calc = [
    "Inactivity > 60s: +30 points",
    "Back navigation > 3: +25 points",
    "Price checks > 5: +20 points",
    "High-value event: +15 points",
    "First-time user: +10 points"
]
for item in friction_calc:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph("B. Channel Router Applies Cost Rules:", style='Intense Quote')
channel_table = doc.add_table(rows=4, cols=3)
channel_table.style = 'Table Grid'
ch_headers = ["Cart Value", "Allowed Channels", "Priority"]
for i, header in enumerate(ch_headers):
    channel_table.rows[0].cells[i].text = header
    channel_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

channel_data = [
    ("< ₹5,000", "None (self-serve only)", "-"),
    ("₹5,000 - ₹25,000", "Chat only", "Normal"),
    ("> ₹25,000", "Chat + Call", "High")
]
for i, (value, channels, priority) in enumerate(channel_data, 1):
    channel_table.rows[i].cells[0].text = value
    channel_table.rows[i].cells[1].text = channels
    channel_table.rows[i].cells[2].text = priority

doc.add_paragraph()
doc.add_paragraph("C. Incidence Created:", style='Intense Quote')
inc_items = [
    "New record created in incidences table",
    "Links to user's context and session",
    "Outcome set to IN_PROGRESS"
]
for item in inc_items:
    doc.add_paragraph(item, style='List Bullet')

# Step 4
doc.add_heading("Step 4: Chat Opens via Freshchat", level=3)
doc.add_paragraph(
    "The agent dashboard shows full user context instantly — no need to ask \"What are you looking at?\""
)
doc.add_paragraph("Agent sees:", style='Intense Quote')
agent_context = [
    "User: John Doe (Gold Tier)",
    "Screen: Checkout",
    "Cart: ₹25,000 | 100 guests",
    "Event: Wedding - Feb 15",
    "Friction Score: 55 (HIGH)",
    "Items: Paneer Tikka, Dal Makhani, Biryani"
]
for item in agent_context:
    doc.add_paragraph(item, style='List Bullet')

# Step 5
doc.add_heading("Step 5: Conversation Happens", level=3)
step5_items = [
    "Every message triggers a webhook to your backend",
    "message_create event → logged to incidence_timeline",
    "Agent helps user complete order with full context visibility"
]
for item in step5_items:
    doc.add_paragraph(item, style='List Bullet')

# Step 6
doc.add_heading("Step 6: Resolution", level=3)
doc.add_paragraph("When chat is resolved:")
step6_items = [
    "Webhook: conversation_resolution is received",
    "Backend updates incidence: outcome → CONVERTED, order_impact → PLACED",
    "Time to resolve is calculated and logged",
    "Analytics Service logs for KPI tracking"
]
for item in step6_items:
    doc.add_paragraph(item, style='List Bullet')

# Flow Summary
doc.add_heading("2.2 Flow Summary Diagram", level=2)
doc.add_paragraph(
    "User Opens App → Browses → Clicks 'Help' → Backend calculates (Friction Score, Allowed Channels, Creates Incidence) "
    "→ Freshchat shows chat with full user context to agent → Webhooks log everything back to database → Resolution → Analytics"
)

doc.add_page_break()



# 3. System Architecture
doc.add_heading("3. System Architecture", level=1)

doc.add_heading("3.1 High-Level Architecture", level=2)
doc.add_paragraph(
    "The system consists of three main layers:\n\n"
    "1. MOBILE APP: React Native application with Freshchat SDK, Context Tracker, and Friction Detector\n\n"
    "2. FRESHCHAT CLOUD: Agent Inbox, User Context storage, Conversation handling\n\n"
    "3. BACKEND SERVICES: FastAPI server with Webhook Handler, Incidence Service, Channel Router, and Analytics Service"
)

doc.add_heading("3.2 Component Overview", level=2)
comp_table = doc.add_table(rows=6, cols=3)
comp_table.style = 'Table Grid'
headers = ["Component", "Technology", "Purpose"]
for i, header in enumerate(headers):
    comp_table.rows[0].cells[i].text = header
    comp_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

components = [
    ("Mobile App", "React Native", "User-facing app with Freshchat SDK"),
    ("API Gateway", "FastAPI (Python)", "REST API, webhook handling"),
    ("Database", "PostgreSQL", "Persistent storage"),
    ("Cache", "Redis", "Real-time context, session data"),
    ("Queue", "Celery + Redis", "Async webhook processing")
]
for i, (comp, tech, purpose) in enumerate(components, 1):
    comp_table.rows[i].cells[0].text = comp
    comp_table.rows[i].cells[1].text = tech
    comp_table.rows[i].cells[2].text = purpose

doc.add_page_break()

# 4. Data Models
doc.add_heading("4. Data Models", level=1)

doc.add_heading("4.1 Incidence Model", level=2)
doc.add_paragraph(
    "The Incidence is the core entity representing every support interaction. "
    "It replaces the traditional 'lead' concept with rich contextual data."
)

doc.add_paragraph("Key fields:", style='Intense Quote')
incidence_fields = [
    ("id", "UUID - Unique identifier"),
    ("user_id", "Reference to user"),
    ("conversation_id", "Freshchat conversation ID"),
    ("stage", "PRE_ORDER or POST_ORDER"),
    ("channel", "IN_APP_CHAT, CALL, or WHATSAPP"),
    ("trigger", "USER_INITIATED or SYSTEM_INITIATED"),
    ("app_screen", "Screen user was on when help requested"),
    ("cart_value", "Cart value at time of incidence"),
    ("friction_score", "Calculated friction score (0-100)"),
    ("outcome", "RESOLVED, DROPPED, CONVERTED, or IN_PROGRESS"),
    ("issue_category", "Categorized issue type"),
    ("time_to_resolve_seconds", "Resolution time")
]

fields_table = doc.add_table(rows=len(incidence_fields)+1, cols=2)
fields_table.style = 'Table Grid'
fields_table.rows[0].cells[0].text = "Field"
fields_table.rows[0].cells[1].text = "Description"
for cell in fields_table.rows[0].cells:
    cell.paragraphs[0].runs[0].font.bold = True

for i, (field, desc) in enumerate(incidence_fields, 1):
    fields_table.rows[i].cells[0].text = field
    fields_table.rows[i].cells[1].text = desc

doc.add_heading("4.2 User Context Model", level=2)
doc.add_paragraph(
    "Real-time user state visible to agents. Stored in Redis for fast access. "
    "Updated on every screen navigation and cart change."
)

context_fields = [
    "User profile (name, email, phone, tier, past orders)",
    "Current session (screen, duration, session start)",
    "Cart state (items, value, guest count, event date)",
    "Behavior signals (friction score, back navigations, payment retries)"
]
for field in context_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_page_break()

# 5. API Specifications
doc.add_heading("5. API Specifications", level=1)

doc.add_heading("5.1 Webhook Handler", level=2)
doc.add_paragraph("Endpoint: POST /webhooks/freshchat", style='Intense Quote')
doc.add_paragraph(
    "Receives real-time events from Freshchat and processes them to create/update Incidences."
)

doc.add_paragraph("Supported Events:")
webhook_events = [
    ("message_create", "Create incidence, log to timeline"),
    ("conversation_assignment", "Update agent assignment"),
    ("conversation_resolution", "Close incidence with outcome"),
    ("conversation_reopen", "Reopen incidence")
]

events_table = doc.add_table(rows=len(webhook_events)+1, cols=2)
events_table.style = 'Table Grid'
events_table.rows[0].cells[0].text = "Event"
events_table.rows[0].cells[1].text = "Action"
for cell in events_table.rows[0].cells:
    cell.paragraphs[0].runs[0].font.bold = True

for i, (event, action) in enumerate(webhook_events, 1):
    events_table.rows[i].cells[0].text = event
    events_table.rows[i].cells[1].text = action

doc.add_heading("5.2 Context API", level=2)
doc.add_paragraph("Base Path: /api/v1/context", style='Intense Quote')

context_endpoints = [
    ("PUT /{user_id}/screen", "Update user's current screen"),
    ("PUT /{user_id}/cart", "Update cart state"),
    ("POST /{user_id}/friction-signal", "Log friction signal")
]

ctx_table = doc.add_table(rows=len(context_endpoints)+1, cols=2)
ctx_table.style = 'Table Grid'
ctx_table.rows[0].cells[0].text = "Endpoint"
ctx_table.rows[0].cells[1].text = "Purpose"
for cell in ctx_table.rows[0].cells:
    cell.paragraphs[0].runs[0].font.bold = True

for i, (endpoint, purpose) in enumerate(context_endpoints, 1):
    ctx_table.rows[i].cells[0].text = endpoint
    ctx_table.rows[i].cells[1].text = purpose

doc.add_heading("5.3 Channel Router API", level=2)
doc.add_paragraph("Endpoint: POST /api/v1/channel/route", style='Intense Quote')
doc.add_paragraph("Business Rules:")
channel_rules = [
    "< ₹5,000: No human support (self-serve only)",
    "₹5,000 - ₹25,000: Chat only",
    "> ₹25,000: Chat + Call",
    "High-importance events (Wedding, Corporate, Religious): Always Chat + Call"
]
for rule in channel_rules:
    doc.add_paragraph(rule, style='List Bullet')

doc.add_page_break()

# 6. Service Layer
doc.add_heading("6. Service Layer", level=1)

doc.add_heading("6.1 Incidence Service", level=2)
doc.add_paragraph("Core service for managing support incidences.")
doc.add_paragraph("Key Methods:")
incidence_methods = [
    ("create()", "Create new incidence with context"),
    ("close()", "Close incidence with resolution details"),
    ("get_by_conversation()", "Lookup by Freshchat conversation ID"),
    ("log_timeline()", "Add event to incidence timeline")
]
for method, desc in incidence_methods:
    doc.add_paragraph(f"{method} — {desc}", style='List Bullet')

doc.add_heading("6.2 Friction Detection Service", level=2)
doc.add_paragraph(
    "Monitors user behavior and calculates friction score. "
    "Triggers help prompts when friction exceeds threshold (60)."
)

doc.add_paragraph("Friction Weights:")
friction_weights = [
    ("Inactivity on checkout (>60s)", "30 points"),
    ("Back navigation loops (>3)", "25 points"),
    ("Excessive price checking (>5)", "20 points"),
    ("High-importance event", "15 points"),
    ("First-time user", "10 points"),
    ("Payment failure", "40 points")
]

friction_table = doc.add_table(rows=len(friction_weights)+1, cols=2)
friction_table.style = 'Table Grid'
friction_table.rows[0].cells[0].text = "Signal"
friction_table.rows[0].cells[1].text = "Weight"
for cell in friction_table.rows[0].cells:
    cell.paragraphs[0].runs[0].font.bold = True

for i, (signal, weight) in enumerate(friction_weights, 1):
    friction_table.rows[i].cells[0].text = signal
    friction_table.rows[i].cells[1].text = weight

doc.add_heading("6.3 Channel Router Service", level=2)
doc.add_paragraph(
    "Routes users to appropriate support channels based on order value and event type. "
    "Enforces cost boundaries to optimize support spend."
)

doc.add_page_break()

# 7. Mobile Integration
doc.add_heading("7. Mobile Integration", level=1)

doc.add_heading("7.1 Freshchat SDK Setup", level=2)
doc.add_paragraph("Key integration steps:")
sdk_steps = [
    "Initialize Freshchat SDK with app credentials on app startup",
    "Identify user after login (set name, email, phone)",
    "Set custom properties for agent context (tier, past orders, etc.)",
    "Sync real-time context on screen navigation and cart updates",
    "Control chat visibility based on channel routing rules"
]
for i, step in enumerate(sdk_steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("7.2 Context Tracker", level=2)
doc.add_paragraph(
    "Tracks user behavior and maintains real-time context. "
    "Syncs with both Freshchat (for agent visibility) and backend (for analytics)."
)
doc.add_paragraph("Tracked behaviors:")
behaviors = [
    "Screen navigation and time spent",
    "Cart updates (value, items, guest count)",
    "Back navigation patterns",
    "Payment attempts and failures",
    "Friction score calculation"
]
for behavior in behaviors:
    doc.add_paragraph(behavior, style='List Bullet')

doc.add_page_break()

# 8. Analytics
doc.add_heading("8. Analytics & Reporting", level=1)

doc.add_heading("8.1 Key Metrics", level=2)
metrics = [
    ("Self-serve conversion rate", "Orders completed without human help / Total orders"),
    ("Assisted conversion rate", "Orders with help that converted / Total assisted"),
    ("Average time to resolve", "Mean resolution time across all incidences"),
    ("Cost per assisted order", "Total support cost / Assisted orders"),
    ("Top friction reasons", "Most common issue categories")
]

metrics_table = doc.add_table(rows=len(metrics)+1, cols=2)
metrics_table.style = 'Table Grid'
metrics_table.rows[0].cells[0].text = "Metric"
metrics_table.rows[0].cells[1].text = "Calculation"
for cell in metrics_table.rows[0].cells:
    cell.paragraphs[0].runs[0].font.bold = True

for i, (metric, calc) in enumerate(metrics, 1):
    metrics_table.rows[i].cells[0].text = metric
    metrics_table.rows[i].cells[1].text = calc

doc.add_heading("8.2 Weekly Report", level=2)
doc.add_paragraph(
    "Generated every Monday, includes: total orders, self-serve rate, "
    "assisted orders, average resolution time, cost analysis, and top 10 friction reasons."
)

doc.add_page_break()

# 9. Deployment
doc.add_heading("9. Deployment", level=1)

doc.add_heading("9.1 Infrastructure Requirements", level=2)
infra = [
    ("API Server", "2 vCPU, 4GB RAM", "~₹3,000/mo"),
    ("PostgreSQL", "2 vCPU, 4GB RAM, 50GB", "~₹2,500/mo"),
    ("Redis", "1GB", "~₹1,000/mo"),
    ("Celery Workers", "1 vCPU, 2GB RAM", "~₹1,500/mo"),
    ("TOTAL", "", "~₹8,000/mo")
]

infra_table = doc.add_table(rows=len(infra)+1, cols=3)
infra_table.style = 'Table Grid'
headers = ["Component", "Specification", "Monthly Cost"]
for i, header in enumerate(headers):
    infra_table.rows[0].cells[i].text = header
    infra_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for i, (comp, spec, cost) in enumerate(infra, 1):
    infra_table.rows[i].cells[0].text = comp
    infra_table.rows[i].cells[1].text = spec
    infra_table.rows[i].cells[2].text = cost

doc.add_heading("9.2 Total Monthly Cost", level=2)
cost_summary = [
    ("Freshchat (6 agents)", "₹9,600"),
    ("Infrastructure", "₹8,000"),
    ("TOTAL", "₹17,600")
]
for item, cost in cost_summary:
    doc.add_paragraph(f"{item}: {cost}")

doc.add_page_break()

# 10. Timeline
doc.add_heading("10. Implementation Timeline", level=1)

timeline = [
    ("Week 1", "Freshchat setup, SDK integration, context passing", "Chat working in app"),
    ("Week 2", "Webhook handler, Incidence Service, database", "Incidences being logged"),
    ("Week 3", "Channel Router, Friction Detection", "Smart routing active"),
    ("Week 4", "Analytics Service, Agent Dashboard config", "Weekly reports")
]

timeline_table = doc.add_table(rows=len(timeline)+1, cols=3)
timeline_table.style = 'Table Grid'
headers = ["Week", "Tasks", "Deliverable"]
for i, header in enumerate(headers):
    timeline_table.rows[0].cells[i].text = header
    timeline_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for i, (week, tasks, deliverable) in enumerate(timeline, 1):
    timeline_table.rows[i].cells[0].text = week
    timeline_table.rows[i].cells[1].text = tasks
    timeline_table.rows[i].cells[2].text = deliverable

doc.add_page_break()

# 11. Success Metrics
doc.add_heading("11. Success Metrics", level=1)

targets = [
    ("Self-serve conversion", "60%", "75%"),
    ("Avg. time to resolve", "<10 min", "<5 min"),
    ("Cost per assisted order", "<₹50", "<₹30"),
    ("Repeat friction issues", "Track baseline", "-30%")
]

targets_table = doc.add_table(rows=len(targets)+1, cols=3)
targets_table.style = 'Table Grid'
headers = ["Metric", "Target (Month 1)", "Target (Month 3)"]
for i, header in enumerate(headers):
    targets_table.rows[0].cells[i].text = header
    targets_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

for i, (metric, m1, m3) in enumerate(targets, 1):
    targets_table.rows[i].cells[0].text = metric
    targets_table.rows[i].cells[1].text = m1
    targets_table.rows[i].cells[2].text = m3

doc.add_paragraph()
doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run("Document Status: Ready for Engineering Review").italic = True
doc.add_paragraph()
next_steps = doc.add_paragraph()
next_steps.alignment = WD_ALIGN_PARAGRAPH.CENTER
next_steps.add_run("Next Steps: Review with team → Sprint planning → Implementation")

# Save document
output_path = r"C:\Users\Syed Ashfaque Hussai\OneDrive\Desktop\Customer Support\Technical_PRD.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
